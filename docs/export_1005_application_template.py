#!/usr/bin/env python3
"""Build the OpenNutri 1005 application in the official Word template."""

from __future__ import annotations

import argparse
import re
import shutil
import subprocess
import tempfile
from copy import deepcopy
from pathlib import Path
from typing import Iterable, List, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_TEMPLATE = ROOT / "1005_basvuru_formu (7).doc"
DEFAULT_EK1_TEMPLATE = ROOT / "1005_basvuru_formu_ek1_kaynaklar_31.07.2018 (4).doc"
DEFAULT_EK2_TEMPLATE = ROOT / "ek-2_butce_ve_gerekcesi_tablosu_1005.docx"
DEFAULT_PROPOSAL = ROOT / "OpenNutriLatestVersion.md"
DEFAULT_BUDGET = ROOT / "OpenNutri_EK2_Butce.md"
DEFAULT_OUTPUT = ROOT / "docs" / "OpenNutri_1005_Application_Filled_Template.docx"
DEFAULT_EK1_OUTPUT = ROOT / "docs" / "OpenNutri_1005_EK1_Kaynaklar.docx"
DEFAULT_EK2_OUTPUT = ROOT / "docs" / "OpenNutri_1005_EK2_Butce_ve_Gerekcesi.docx"


MAIN_SECTION_TITLES = {
    "1": "ULUSAL KAZANIM ve PROJENİN ÖNEMİ",
    "2": "AMAÇ VE HEDEFLER",
    "3": "PATENT, FAYDALI MODEL, TESCİL TARAMASI, YENİLİKÇİ YÖNÜ VE TİCARİLEŞME POTANSİYELİ",
    "4": "YÖNTEM",
    "5": "PROJE YÖNETİMİ",
    "6": "YAYGIN ETKİ",
    "EK-1": "KAYNAKLAR",
}


def convert_template_to_docx(template: Path, workdir: Path) -> Path:
    if not template.exists():
        raise FileNotFoundError(f"Template not found: {template}")
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise RuntimeError("LibreOffice/soffice is required to convert the .doc template.")
    subprocess.run(
        [soffice, "--headless", "--convert-to", "docx", "--outdir", str(workdir), str(template)],
        cwd=ROOT,
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )
    converted = workdir / f"{template.stem}.docx"
    if not converted.exists():
        matches = list(workdir.glob("*.docx"))
        if not matches:
            raise RuntimeError("LibreOffice reported success but did not create a DOCX.")
        converted = matches[0]
    return converted


def split_top_sections(markdown: str) -> dict[str, str]:
    sections: dict[str, list[str]] = {}
    current = "cover"
    sections[current] = []
    for line in markdown.splitlines():
        match = re.match(r"^#\s+((?:\d+|EK-1)\b)(.*)$", line)
        if match:
            current = match.group(1)
            sections[current] = []
            continue
        sections.setdefault(current, []).append(line)
    return {key: "\n".join(value).strip() for key, value in sections.items()}


def extract_cover(markdown: str) -> dict[str, str]:
    fields = {
        "Proje Başlığı": "",
        "Proje Yürütücüsü": "",
        "Projenin Yürütüleceği Kurum/Kuruluş": "",
    }
    for label in fields:
        pattern = re.compile(rf"^\*\*{re.escape(label)}:\*\*\s*(.+)$", re.MULTILINE)
        match = pattern.search(markdown)
        if match:
            fields[label] = match.group(1).strip()
    return fields


def remove_newpage(markdown: str) -> str:
    return "\n".join(line for line in markdown.splitlines() if line.strip() != r"\newpage").strip()


def section_without_subheading(markdown: str, heading_regex: str) -> str:
    lines = markdown.splitlines()
    out: list[str] = []
    skip_first_heading = True
    for line in lines:
        if skip_first_heading and re.match(heading_regex, line):
            skip_first_heading = False
            continue
        out.append(line)
    return "\n".join(out).strip()


def parse_markdown_table(lines: Sequence[str], start: int) -> tuple[list[list[str]], int] | None:
    if start + 1 >= len(lines):
        return None
    first = lines[start].strip()
    second = lines[start + 1].strip()
    if not (first.startswith("|") and first.endswith("|") and second.startswith("|")):
        return None
    separator_cells = [cell.strip() for cell in second.strip("|").split("|")]
    if not separator_cells or not all(re.fullmatch(r":?-{3,}:?", cell) for cell in separator_cells):
        return None
    rows: list[list[str]] = []
    index = start
    while index < len(lines):
        line = lines[index].strip()
        if not (line.startswith("|") and line.endswith("|")):
            break
        if index != start + 1:
            rows.append([cell.strip() for cell in line.strip("|").split("|")])
        index += 1
    return rows, index


def first_markdown_table(markdown: str) -> list[list[str]]:
    lines = markdown.splitlines()
    for i in range(len(lines)):
        parsed = parse_markdown_table(lines, i)
        if parsed:
            return parsed[0]
    return []


def markdown_without_first_table(markdown: str) -> str:
    lines = markdown.splitlines()
    for i in range(len(lines)):
        parsed = parse_markdown_table(lines, i)
        if parsed:
            _, end = parsed
            return "\n".join(lines[:i] + lines[end:]).strip()
    return markdown.strip()


def subsection(markdown: str, number: str) -> str:
    pattern = re.compile(rf"^##\s+{re.escape(number)}\b.*$", re.MULTILINE)
    match = pattern.search(markdown)
    if not match:
        return ""
    start = match.start()
    next_match = re.search(r"^##\s+\d+(?:\.\d+)?\b.*$", markdown[match.end() :], re.MULTILINE)
    end = match.end() + next_match.start() if next_match else len(markdown)
    return markdown[start:end].strip()


def strip_first_heading(markdown: str) -> str:
    lines = markdown.splitlines()
    if lines and lines[0].lstrip().startswith("#"):
        return "\n".join(lines[1:]).strip()
    return markdown.strip()


def split_gantt(section_51: str) -> tuple[str, str]:
    marker = "**İş-Zaman Çizelgesi"
    if marker not in section_51:
        return section_51, ""
    table_part, gantt_part = section_51.split(marker, 1)
    return table_part.strip(), (marker + gantt_part).strip()


def clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def clear_cell(cell: _Cell) -> None:
    tc = cell._tc
    for child in list(tc):
        if child.tag != qn("w:tcPr"):
            tc.remove(child)
    cell.add_paragraph()


def ensure_cell_paragraph(cell: _Cell) -> Paragraph:
    if not cell.paragraphs:
        return cell.add_paragraph()
    return cell.paragraphs[0]


def set_run_font(run, *, size: float = 9, bold: bool | None = None, italic: bool | None = None, mono: bool = False) -> None:
    run.font.name = "Courier New" if mono else "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New" if mono else "Arial")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


INLINE_TOKEN = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def add_inline(paragraph: Paragraph, text: str, *, size: float = 9) -> None:
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = text.replace("\\≥", "≥").replace("\\<", "<").replace("\\%", "%")
    pos = 0
    for match in INLINE_TOKEN.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, italic=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, mono=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size)


def format_paragraph(paragraph: Paragraph, *, size: float = 9, bold: bool = False, italic: bool = False) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold if bold else None, italic=italic if italic else None)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0


def add_markdown_paragraph(cell: _Cell, text: str, *, kind: str = "body") -> Paragraph:
    paragraph = cell.add_paragraph()
    if kind == "h2":
        add_inline(paragraph, text, size=9.5)
        format_paragraph(paragraph, size=9.5, bold=True)
    elif kind == "h3":
        add_inline(paragraph, text, size=9)
        format_paragraph(paragraph, size=9, bold=True)
    elif kind == "quote":
        add_inline(paragraph, text, size=8.5)
        format_paragraph(paragraph, size=8.5, italic=True)
    elif kind == "code":
        run = paragraph.add_run(text)
        set_run_font(run, size=7.5, mono=True)
        format_paragraph(paragraph, size=7.5)
    else:
        add_inline(paragraph, text, size=9)
        format_paragraph(paragraph, size=9)
    return paragraph


def append_docx_table(cell: _Cell, rows: list[list[str]], *, font_size: float = 8) -> Table:
    if not rows:
        return cell.add_table(rows=1, cols=1)
    table = cell.add_table(rows=len(rows), cols=max(len(row) for row in rows))
    try:
        table.style = "Table Grid"
    except KeyError:
        # Older localized templates may not carry the English table style name.
        pass
    set_explicit_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_index, row in enumerate(rows):
        for c_index in range(len(table.columns)):
            target = table.cell(r_index, c_index)
            clear_cell(target)
            if c_index < len(row):
                add_markdown_to_cell(target, row[c_index], clear=False, compact=True, font_size=font_size)
            for paragraph in target.paragraphs:
                format_paragraph(paragraph, size=font_size, bold=(r_index == 0))
            target.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    return table


def set_explicit_table_borders(table: Table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def add_markdown_to_cell(
    cell: _Cell,
    markdown: str,
    *,
    clear: bool = True,
    compact: bool = False,
    font_size: float = 9,
    allow_tables: bool = True,
) -> None:
    if clear:
        clear_cell(cell)
        # Remove the placeholder paragraph inserted by clear_cell(). Otherwise,
        # sections that begin with a heading/table can have the next body
        # paragraph back-filled above that first block.
        for child in list(cell._tc):
            if child.tag != qn("w:tcPr"):
                cell._tc.remove(child)
    first_para_used = True

    lines = remove_newpage(markdown).splitlines()
    paragraph_buffer: list[str] = []
    in_code = False
    code_lines: list[str] = []
    i = 0

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer, first_para_used
        if not paragraph_buffer:
            return
        text = " ".join(part.strip() for part in paragraph_buffer if part.strip())
        paragraph_buffer = []
        if not text:
            return
        if not first_para_used and cell.paragraphs:
            paragraph = cell.paragraphs[0]
            first_para_used = True
            add_inline(paragraph, text, size=font_size)
            format_paragraph(paragraph, size=font_size)
        else:
            paragraph = add_markdown_paragraph(cell, text)
            format_paragraph(paragraph, size=font_size)

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                for code_line in code_lines:
                    add_markdown_paragraph(cell, code_line, kind="code")
                code_lines = []
                in_code = False
            else:
                flush_paragraph()
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        parsed_table = parse_markdown_table(lines, i)
        if parsed_table:
            flush_paragraph()
            rows, end = parsed_table
            if allow_tables:
                append_docx_table(cell, rows, font_size=7.5 if not compact else 7)
            else:
                table_header = rows[0] if rows else []
                for row in rows[1:] if len(rows) > 1 else rows:
                    parts = []
                    for c_index, value in enumerate(row):
                        if not value.strip():
                            continue
                        label = table_header[c_index].strip() if c_index < len(table_header) else ""
                        if label:
                            parts.append(f"{label}: {value}")
                        else:
                            parts.append(value)
                    if parts:
                        paragraph = cell.add_paragraph()
                        paragraph.paragraph_format.left_indent = Cm(0.35)
                        paragraph.paragraph_format.first_line_indent = Cm(-0.2)
                        run = paragraph.add_run("• ")
                        set_run_font(run, size=font_size)
                        add_inline(paragraph, "; ".join(parts), size=font_size)
                        format_paragraph(paragraph, size=font_size)
            i = end
            continue

        if not stripped or stripped == "---":
            flush_paragraph()
            i += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            level = len(heading.group(1))
            text = heading.group(2).strip()
            add_markdown_paragraph(cell, text, kind="h2" if level <= 2 else "h3")
            i += 1
            continue

        if stripped.startswith(">"):
            flush_paragraph()
            quote = stripped.lstrip(">").strip()
            add_markdown_paragraph(cell, quote, kind="quote")
            i += 1
            continue

        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if bullet or numbered:
            flush_paragraph()
            paragraph = cell.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.35)
            paragraph.paragraph_format.first_line_indent = Cm(-0.2)
            prefix = "• " if bullet else f"{numbered.group(1)}. "
            run = paragraph.add_run(prefix)
            set_run_font(run, size=font_size)
            add_inline(paragraph, bullet.group(1) if bullet else numbered.group(2), size=font_size)
            format_paragraph(paragraph, size=font_size)
            i += 1
            continue

        paragraph_buffer.append(stripped)
        i += 1

    flush_paragraph()
    if in_code:
        for code_line in code_lines:
            add_markdown_paragraph(cell, code_line, kind="code")
    if not cell.paragraphs and not cell.tables:
        cell.add_paragraph()


def fill_cover_table(table: Table, fields: dict[str, str]) -> None:
    for row, (label, value) in zip(table.rows, fields.items()):
        cell = row.cells[0]
        clear_cell(cell)
        paragraph = ensure_cell_paragraph(cell)
        run = paragraph.add_run(f"{label}: ")
        set_run_font(run, size=9, bold=True)
        add_inline(paragraph, value, size=9)
        format_paragraph(paragraph, size=9)


def remove_rows_after_header(table: Table) -> None:
    for row in list(table.rows)[1:]:
        table._tbl.remove(row._tr)


def set_table_rows(table: Table, rows: list[list[str]], *, font_size: float = 7.5, keep_header: bool = True) -> None:
    if not rows:
        return
    header = rows[0]
    data = rows[1:] if keep_header else rows
    remove_rows_after_header(table)
    for c_index, text in enumerate(header[: len(table.columns)]):
        cell = table.cell(0, c_index)
        clear_cell(cell)
        add_markdown_to_cell(cell, text, clear=False, compact=True, font_size=font_size)
        for paragraph in cell.paragraphs:
            format_paragraph(paragraph, size=font_size, bold=True)
    for row_data in data:
        row = table.add_row()
        for c_index, cell in enumerate(row.cells):
            clear_cell(cell)
            text = row_data[c_index] if c_index < len(row_data) else ""
            add_markdown_to_cell(cell, text, clear=False, compact=True, font_size=font_size)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_merged_row(table: Table, markdown: str, *, font_size: float = 7.5) -> None:
    row = table.add_row()
    merged = row.cells[0]
    for cell in row.cells[1:]:
        merged = merged.merge(cell)
    clear_cell(merged)
    add_markdown_to_cell(merged, markdown, clear=False, compact=True, font_size=font_size)


def insert_paragraph_before_table(table: Table, text: str) -> Paragraph:
    paragraph_element = OxmlElement("w:p")
    table._tbl.addprevious(paragraph_element)
    paragraph = Paragraph(paragraph_element, table._parent)
    add_inline(paragraph, text, size=9)
    format_paragraph(paragraph, size=9)
    return paragraph


def insert_markdown_after_paragraph(paragraph: Paragraph, markdown: str) -> None:
    holder = Document()
    cell = holder.add_table(rows=1, cols=1).cell(0, 0)
    add_markdown_to_cell(cell, markdown, clear=True)
    elements = [deepcopy(child) for child in cell._tc if child.tag in {qn("w:p"), qn("w:tbl")}]
    anchor = paragraph._p
    for element in reversed(elements):
        anchor.addnext(element)


def append_markdown_to_document(doc: Document, markdown: str) -> None:
    holder = Document()
    cell = holder.add_table(rows=1, cols=1).cell(0, 0)
    add_markdown_to_cell(cell, markdown, clear=True)
    body = doc.element.body
    sect_pr = body.sectPr
    insert_index = list(body).index(sect_pr) if sect_pr is not None else len(body)
    for child in cell._tc:
        if child.tag not in {qn("w:p"), qn("w:tbl")}:
            continue
        body.insert(insert_index, deepcopy(child))
        insert_index += 1


def remove_body_after_first_heading(doc: Document) -> None:
    body = doc.element.body
    non_empty_seen = False
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        text = ""
        if child.tag == qn("w:p"):
            text = "".join(node.text or "" for node in child.iter(qn("w:t"))).strip()
        if not non_empty_seen and text:
            non_empty_seen = True
            continue
        if non_empty_seen:
            body.remove(child)


def set_cell_markdown(cell: _Cell, markdown: str, *, font_size: float = 8, bold: bool = False) -> None:
    clear_cell(cell)
    add_markdown_to_cell(cell, markdown, clear=False, compact=True, font_size=font_size, allow_tables=False)
    for paragraph in cell.paragraphs:
        format_paragraph(paragraph, size=font_size, bold=bold)


def set_table_row_values(
    table: Table,
    row_index: int,
    values: Sequence[str],
    *,
    font_size: float = 8,
    bold_columns: set[int] | None = None,
) -> None:
    bold_columns = bold_columns or set()
    for col_index, value in enumerate(values):
        if col_index >= len(table.columns):
            break
        set_cell_markdown(
            table.cell(row_index, col_index),
            value,
            font_size=font_size,
            bold=col_index in bold_columns,
        )


def style_document(doc: Document) -> None:
    for style_name in ["Normal", "Body Text"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "Arial"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
            style.font.size = Pt(9)
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            for run in paragraph.runs:
                if run.font.name is None:
                    set_run_font(run, size=9)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        if run.font.name is None:
                            set_run_font(run, size=8.5)


def find_paragraph(doc: Document, exact_text: str) -> Paragraph | None:
    for paragraph in doc.paragraphs:
        if " ".join(paragraph.text.split()) == exact_text:
            return paragraph
    return None


def build_application(
    template: Path,
    proposal: Path,
    budget: Path,
    output: Path,
    *,
    include_references_annex: bool = False,
    include_budget_annex: bool = False,
) -> None:
    proposal_text = proposal.read_text(encoding="utf-8")
    budget_text = budget.read_text(encoding="utf-8") if budget.exists() else ""
    sections = split_top_sections(proposal_text)

    with tempfile.TemporaryDirectory(prefix="opennutri_1005_") as tmp:
        converted = convert_template_to_docx(template, Path(tmp))
        doc = Document(converted)

    fill_cover_table(doc.tables[0], extract_cover(proposal_text))

    for table_index, section_key in [(1, "1"), (2, "2"), (3, "3"), (4, "4")]:
        add_markdown_to_cell(
            doc.tables[table_index].cell(0, 0),
            sections.get(section_key, ""),
            allow_tables=False,
        )

    section_5 = sections.get("5", "")
    section_51 = strip_first_heading(subsection(section_5, "5.1"))
    wp_table_part, gantt = split_gantt(section_51)
    set_table_rows(doc.tables[5], first_markdown_table(wp_table_part), font_size=9)
    if gantt:
        add_merged_row(doc.tables[5], gantt, font_size=9)

    section_52 = strip_first_heading(subsection(section_5, "5.2"))
    set_table_rows(doc.tables[6], first_markdown_table(section_52), font_size=9)

    section_53 = strip_first_heading(subsection(section_5, "5.3"))
    set_table_rows(doc.tables[7], first_markdown_table(section_53), font_size=9)

    section_6 = sections.get("6", "")
    section_61 = strip_first_heading(subsection(section_6, "6.1"))
    set_table_rows(doc.tables[8], first_markdown_table(section_61), font_size=9)

    section_62 = strip_first_heading(subsection(section_6, "6.2"))
    impact_intro = markdown_without_first_table(section_62).strip()
    if impact_intro:
        insert_paragraph_before_table(doc.tables[9], impact_intro)
    set_table_rows(doc.tables[9], first_markdown_table(section_62), font_size=9)

    section_63 = strip_first_heading(subsection(section_6, "6.3"))
    add_markdown_to_cell(doc.tables[10].cell(0, 0), section_63, allow_tables=False)

    section_64 = strip_first_heading(subsection(section_6, "6.4"))
    set_table_rows(doc.tables[11], first_markdown_table(section_64), font_size=9)

    add_markdown_to_cell(doc.tables[12].cell(0, 0), "Ek bilgi yoktur.")

    ek1_paragraph = find_paragraph(doc, "EK-1: KAYNAKLAR")
    if include_references_annex and ek1_paragraph is not None and sections.get("EK-1"):
        insert_markdown_after_paragraph(ek1_paragraph, sections["EK-1"])

    ek2_paragraph = find_paragraph(doc, "EK-2: BÜTÇE VE GEREKÇESİ")
    if include_budget_annex and ek2_paragraph is not None and budget_text:
        insert_markdown_after_paragraph(ek2_paragraph, strip_first_heading(budget_text))

    style_document(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def build_ek1(template: Path, proposal: Path, output: Path) -> None:
    proposal_text = proposal.read_text(encoding="utf-8")
    references = split_top_sections(proposal_text).get("EK-1", "")
    if not references:
        raise RuntimeError("No EK-1 references section found in proposal markdown.")
    with tempfile.TemporaryDirectory(prefix="opennutri_ek1_") as tmp:
        converted = convert_template_to_docx(template, Path(tmp))
        doc = Document(converted)
    remove_body_after_first_heading(doc)
    append_markdown_to_document(doc, references)
    style_document(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def build_ek2(template: Path, budget: Path, output: Path) -> None:
    if not template.exists():
        raise FileNotFoundError(f"EK-2 template not found: {template}")
    if not budget.exists():
        raise FileNotFoundError(f"Budget markdown not found: {budget}")
    doc = Document(template)

    # General budget table.
    set_table_row_values(
        doc.tables[0],
        1,
        ["TÜBİTAK'tan Talep Edilen Katkı", "242.000", "25.000", "60.000", "0", "45.000", "828.000", "0", "1.200.000"],
        font_size=6.3,
        bold_columns={0, 8},
    )
    set_table_row_values(
        doc.tables[0],
        2,
        ["Öneren Kuruluş Katkısı", "0", "0", "0", "0", "0", "0", "0", "0"],
        font_size=6.3,
        bold_columns={0},
    )
    set_table_row_values(
        doc.tables[0],
        3,
        ["Destekleyen Diğer Kuruluş Katkısı", "0", "0", "0", "0", "0", "0", "0", "0"],
        font_size=6.3,
        bold_columns={0},
    )
    set_table_row_values(
        doc.tables[0],
        4,
        ["TOPLAM", "242.000", "25.000", "60.000", "0", "45.000", "828.000", "0", "1.200.000"],
        font_size=6.3,
        bold_columns={0, 8},
    )

    # Machine/equipment.
    equipment_rows = [
        [
            "GPU iş istasyonu / 1 adet",
            "L3 ağırlıkları açık modellerin geliştirme/testi, çıkarım servisleri, Öğrenilmiş Yönlendirici eğitimi ve PEFT/LoRA-QLoRA ölçeğinde yerel ince ayar.",
            "Tek GPU, 16-24 GB VRAM sınıfı; çok çekirdekli CPU; 64-128 GB RAM; 1-2 TB NVMe. Kesin model proforma aşamasında belirlenecektir.",
            "170.000",
        ],
        [
            "NAS depolama + diskler / 1 adet",
            "Birincil veri tabanı, PDF/tam metin önbelleği, model kontrol noktaları ve yedekleme.",
            "Yedekli NAS, yaklaşık 8-12 TB ham depolama kapasitesi.",
            "55.000",
        ],
        [
            "Kesintisiz güç kaynağı (KGK) + ağ donanımı / 1 set",
            "Kesintisiz operasyon ve veri bütünlüğü.",
            "İş istasyonu ve NAS için uygun kapasiteli KGK; temel ağ bağlantı donanımı.",
            "17.000",
        ],
    ]
    for offset, row in enumerate(equipment_rows, start=2):
        set_table_row_values(doc.tables[1], offset, row, font_size=6.1)

    # Consumables.
    set_table_row_values(
        doc.tables[2],
        2,
        [
            "SSD/HDD yedekleri, bileşenler, kablolar ve küçük donanım",
            "Depolama genişletme, yedek parça ve laboratuvar/altyapı sarf ihtiyaçları.",
            "25.000",
        ],
        font_size=6.5,
    )
    set_table_row_values(doc.tables[2], 4, ["TOPLAM", "", "25.000"], font_size=6.5, bold_columns={0, 2})

    # Services.
    service_rows = [
        [
            "Ticari LLM API kullanımı",
            "Model sağlayıcıları",
            "L4 yükseltme katmanı, maliyet-kalite kıyaslaması ve yalnızca başarısız alt görevlerde sınırlı ticari API çağrıları.",
            "35.000",
        ],
        [
            "Bulut yedekleme / dağıtım",
            "Bulut hizmet sağlayıcıları",
            "Felaket kurtarma, yedekleme ve API dağıtım esnekliği.",
            "10.000",
        ],
        [
            "TTO patent taraması güncelleme + akademik redaksiyon/çeviri",
            "Kurum TTO / dış hizmet sağlayıcı",
            "Patent ön taramasının ürünleşme aşamasında güncellenmesi ve yayın redaksiyonu.",
            "15.000",
        ],
    ]
    for offset, row in enumerate(service_rows, start=2):
        set_table_row_values(doc.tables[3], offset, row, font_size=6.4)
    set_table_row_values(doc.tables[3], 4, service_rows[2], font_size=6.4)

    # Representation/promotion: no requested budget in the current draft.
    set_table_row_values(doc.tables[4], 7, ["Toplam", "", "", "0"], font_size=6.5, bold_columns={0, 3})

    # Field work and field-work travel are not planned.
    set_table_row_values(doc.tables[5], 4, ["TOPLAM", "", "", "", "", "", "", "", "", "0"], font_size=5.5, bold_columns={0, 9})
    set_table_row_values(doc.tables[6], 7, ["TOPLAM (TL)", "", "", "", "", "", "0"], font_size=6.2, bold_columns={0, 6})
    set_cell_markdown(doc.tables[7].cell(0, 0), "Yurt içi saha çalışması planlanmamaktadır.", font_size=8)
    set_table_row_values(doc.tables[8], 4, ["TOPLAM", "", "", "", "", "", "", "", "", "0"], font_size=5.5, bold_columns={0, 9})
    set_table_row_values(doc.tables[9], 7, ["TOPLAM (TL)", "", "", "", "", "", "0"], font_size=6.2, bold_columns={0, 6})

    # Non-field travel.
    set_table_row_values(doc.tables[10], 2, ["Yurt İçi / Yurt Dışı Seyahat", "45.000"], font_size=7, bold_columns={0})
    set_table_row_values(doc.tables[10], 3, ["Yurt Dışı Uçak Bileti (TÜBİTAK'tan talep edilen)", "0"], font_size=7, bold_columns={0})

    # Scholarships, aggregated by level because two food-engineering students move from undergraduate to MSc.
    set_table_row_values(
        doc.tables[11],
        2,
        ["Lisans Öğrencisi", "48", "6.000", "288.000"],
        font_size=7,
    )
    set_table_row_values(
        doc.tables[11],
        3,
        ["Yüksek Lisans Öğrencisi (çalışmayan)", "24", "22.500", "540.000"],
        font_size=7,
    )
    set_table_row_values(doc.tables[11], 4, ["TOPLAM", "", "", "828.000"], font_size=7, bold_columns={0, 3})

    # Temporary workers are not requested.
    set_table_row_values(doc.tables[13], 4, ["TOPLAM", "", "", "", "0"], font_size=7, bold_columns={0, 4})

    style_document(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main(argv: Iterable[str] | None = None) -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--template", type=Path, default=DEFAULT_TEMPLATE)
    parser.add_argument("--proposal", type=Path, default=DEFAULT_PROPOSAL)
    parser.add_argument("--budget", type=Path, default=DEFAULT_BUDGET)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--ek1-template", type=Path, default=DEFAULT_EK1_TEMPLATE)
    parser.add_argument("--ek1-output", type=Path, default=DEFAULT_EK1_OUTPUT)
    parser.add_argument("--ek2-template", type=Path, default=DEFAULT_EK2_TEMPLATE)
    parser.add_argument("--ek2-output", type=Path, default=DEFAULT_EK2_OUTPUT)
    parser.add_argument(
        "--single-file-annexes",
        action="store_true",
        help="Append EK-1/EK-2 content inside the main form instead of producing separate annex files.",
    )
    parser.add_argument(
        "--main-only",
        action="store_true",
        help="Only write the main form; do not generate separate EK-1/EK-2 files.",
    )
    args = parser.parse_args(argv)
    build_application(
        args.template,
        args.proposal,
        args.budget,
        args.output,
        include_references_annex=args.single_file_annexes,
        include_budget_annex=args.single_file_annexes,
    )
    print(f"Created {args.output}")
    if not args.single_file_annexes and not args.main_only:
        build_ek1(args.ek1_template, args.proposal, args.ek1_output)
        print(f"Created {args.ek1_output}")
        build_ek2(args.ek2_template, args.budget, args.ek2_output)
        print(f"Created {args.ek2_output}")


if __name__ == "__main__":
    main()
