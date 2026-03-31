from __future__ import annotations

import re
import subprocess
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH


BASE_DIR = Path(__file__).resolve().parent
MARKDOWN_PATH = BASE_DIR / "3_Sinif_Proje_Raporu_YAZILIM_KTUN_TR.md"
DOCX_PATH = BASE_DIR / "docx" / "3_Sinif_Proje_Raporu_YAZILIM_KTUN_TR.docx"
PDF_DIR = BASE_DIR / "pdf"


def ensure_styles(document: Document) -> None:
    for section in document.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(20)
        section.right_margin = Mm(20)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)

    if "OpenNutriCode" not in document.styles:
        style = document.styles.add_style("OpenNutriCode", 1)
        style.font.name = "Courier New"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        style.font.size = Pt(9)


def add_inline_runs(paragraph: Paragraph, text: str) -> None:
    token_pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    parts = token_pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        else:
            paragraph.add_run(part)


def flush_paragraph(document: Document, buffer: list[str]) -> None:
    if not buffer:
        return
    paragraph = document.add_paragraph()
    add_inline_runs(paragraph, " ".join(item.strip() for item in buffer if item.strip()))
    buffer.clear()


def is_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    return bool(stripped) and all(ch in "|:- " for ch in stripped)


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def render_table(document: Document, lines: list[str]) -> None:
    rows = [split_table_row(line) for line in lines if line.strip()]
    if len(rows) < 2:
        for line in lines:
            paragraph = document.add_paragraph()
            add_inline_runs(paragraph, line)
        return

    header = rows[0]
    body = [row for row in rows[1:] if not is_table_separator("| " + " | ".join(row) + " |")]
    table = document.add_table(rows=1 + len(body), cols=len(header))
    table.style = "Table Grid"

    for idx, value in enumerate(header):
        table.rows[0].cells[idx].text = value

    for row_idx, row in enumerate(body, start=1):
        for col_idx, value in enumerate(row):
            if col_idx < len(table.rows[row_idx].cells):
                table.rows[row_idx].cells[col_idx].text = value


def render_code_block(document: Document, lines: Iterable[str]) -> None:
    for line in lines:
        paragraph = document.add_paragraph(style="OpenNutriCode")
        paragraph.add_run(line.rstrip("\n"))


def render_image(document: Document, markdown_path: Path, image_ref: str) -> None:
    image_path = (markdown_path.parent / image_ref).resolve()
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.2))


def render_markdown(document: Document, markdown_path: Path) -> None:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    buffer: list[str] = []
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped == r"\newpage":
            flush_paragraph(document, buffer)
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            i += 1
            continue

        image_match = re.match(r"^!\[[^\]]*\]\(([^)]+)\)$", stripped)
        if image_match:
            flush_paragraph(document, buffer)
            render_image(document, markdown_path, image_match.group(1))
            i += 1
            continue

        if stripped.startswith("```"):
            flush_paragraph(document, buffer)
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            render_code_block(document, code_lines)
            i += 1
            continue

        if not stripped:
            flush_paragraph(document, buffer)
            i += 1
            continue

        if is_table_line(line):
            flush_paragraph(document, buffer)
            table_lines: list[str] = []
            while i < len(lines) and is_table_line(lines[i]):
                table_lines.append(lines[i])
                i += 1
            render_table(document, table_lines)
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading:
            flush_paragraph(document, buffer)
            level = min(len(heading.group(1)), 4)
            text = heading.group(2).strip()
            paragraph = document.add_heading(level=level)
            add_inline_runs(paragraph, text)
            i += 1
            continue

        ordered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if ordered:
            flush_paragraph(document, buffer)
            paragraph = document.add_paragraph(style="List Number")
            add_inline_runs(paragraph, ordered.group(2).strip())
            i += 1
            continue

        bullet = re.match(r"^-\s+(.*)$", stripped)
        if bullet:
            flush_paragraph(document, buffer)
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_runs(paragraph, bullet.group(1).strip())
            i += 1
            continue

        if stripped.endswith("  "):
            buffer.append(stripped.rstrip())
            flush_paragraph(document, buffer)
            i += 1
            continue

        buffer.append(stripped)
        i += 1

    flush_paragraph(document, buffer)


def export_docx(markdown_path: Path, output_path: Path) -> None:
    document = Document()
    ensure_styles(document)
    render_markdown(document, markdown_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def convert_docx_to_pdf(docx_path: Path, pdf_dir: Path) -> None:
    pdf_dir.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [
            "libreoffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(pdf_dir),
            str(docx_path),
        ],
        check=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )


def main() -> None:
    export_docx(MARKDOWN_PATH, DOCX_PATH)
    convert_docx_to_pdf(DOCX_PATH, PDF_DIR)
    print("Created", DOCX_PATH)
    print("Created", PDF_DIR / f"{DOCX_PATH.stem}.pdf")


if __name__ == "__main__":
    main()
